from docx import Document
import PyPDF2
import io

from services.ai_init import get_genai_client

#raw string -> process
def preProcessDocument(rawContent: str) -> str:
    lines = rawContent.splitlines()
    cleanedLines = []
    for line in lines:
        if line.strip() == '':
            continue
        cleanedLines.append(line.strip())

    cleanContent = ''.join(cleanedLines)

    return cleanContent


# error class - for the file parsing related errors
class FileProcessingError(Exception):
    pass


class getFileContents:
    def __init__(self):
        self.mime_type_dict = {
            "pdf": "application/pdf",
            "csv": "text/csv"
        }

        self.upload_file_prompt = f"""
        Role: Expert in document analysis and RAG (retrieval-augmented generation) optimization  

        Task: A user will upload a document. Your task is to extract, interpret, and return the complete content of the file in plain text, optimized for use in RAG pipelines.  

        Instructions:  
        1. Content Extraction  
        - Accurately extract all the contents from the input.  
        - Ensure the output is plain text (no markdown, bold, italics, or headings except numbering, should be RAG optimized).  
        - The output should be directly usable in RAG pipelines.  

        2. Non-Textual Elements (tables, images, charts, graphs, CSVs, structured data, etc.)  
        - Do not skip any element; retain full context.  
        - Provide a clear descriptive title for each element (RAG-optimized, concise, unambiguous).  
        - Convert structured/visual data into a human-readable text description:  
            - Tables/CSVs → Convert each row into clear, short, human-readable sentences - for RAG. Each key detail should form 
              its own sentence. Example: ‘Product A is priced at 200. It has a quantity of 50
            - Graphs/Charts → Describe the variables, axes, trends, and key values.  
            - Images → Describe the contents with context and relevance to the document.  
        - Always present both the raw extracted text/data and the interpreted explanation.  

        3. Clarity & Consistency  
        - Output must be linear, contextually complete, and self-contained (no references like “see above”).  
        - Use simple, clear sentences so the text is both machine-processable and human-readable.  
        - Do not omit or summarize; include everything, even if approximate (e.g., estimated words in handwritten text).  
        """

    async def basic_pdf(
        self,
        file_content: bytes,
        # file_name: str
    ) -> str:
        """Extract text content from PDF using library"""
        try:
            pdf_file = io.BytesIO(file_content)
            document = PyPDF2.PdfReader(pdf_file)

            docContents = ''
            for page in document.pages:
                docContents += page.extract_text()
            
            finalContents = preProcessDocument(docContents)
            return finalContents

        except Exception as e:
            print(f'Error Processing the PDF (basic): {e}')
            raise FileProcessingError(
                f'Error Processing the PDF (basic): {e}'
            ) from e


    async def using_llm(
        self,
        file_content: bytes,
        # file_name: str,
        content_type: str
    ) -> str:
        """Extract text content from PDF using Gemini (LLM)"""
        try:
            genai_client = get_genai_client()
            pdf_file = io.BytesIO(file_content)

            # upload the file
            mime_type = self.mime_type_dict[content_type]
            upload_file = await genai_client.aio.files.upload(
                file = pdf_file,
                config = dict (
                    mime_type = mime_type
                )
            )

            prompt = self.upload_file_prompt
            response = await genai_client.aio.models.generate_content(
                model = "gemini-2.0-flash",
                contents = [upload_file],
                config = {
                    'system_instruction': prompt
                }
            )

            print("\nLLM file processing input tokens: ", response.usage_metadata.prompt_token_count)
            print("LLM file processing output tokens: ", response.usage_metadata.candidates_token_count, "\n")

            estimated_tokens = response.usage_metadata.prompt_token_count + response.usage_metadata.candidates_token_count

            return response.text

        except Exception as e:
            print(f'Error Processing the PDF (advanced): {e}')
            raise FileProcessingError(
                f'Error Processing the PDF (advanced): {e}'
            ) from e

        
    async def basic_docx(
        self,
        file_content: bytes,
        # file_name: str
    ) -> str:
        """Extract text content from DOCX using library"""
        
        try:
            file = io.BytesIO(file_content)
            doc_file = Document(file)

            full_text = ''
            # paragraph content
            for paragraph in doc_file.paragraphs:
                full_text = full_text + paragraph.text + '\n'
            
            # table content
            for table in doc_file.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text = full_text + ("\t".join(row_text))

            finalContents = preProcessDocument(full_text)
            return finalContents

        except Exception as e:
            print(f'Error Processing the DOCX (basic): {e}')
            raise FileProcessingError(
                f'Error Processing the DOCX (basic): {e}'
            ) from e
        

file_parser = getFileContents()
